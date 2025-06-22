"""
Unit tests for the output_formatter module.

Tests document output functionality for various formats.
"""

import pytest
from unittest.mock import Mock, patch, MagicMock
from pathlib import Path
import tempfile
import shutil
from docx import Document

from video_draft_creator.output_formatter import (
    OutputFormatter,
    create_formatter_from_config
)


class TestOutputFormatter:
    """Test cases for OutputFormatter class."""

    @pytest.fixture
    def formatter(self, sample_config, temp_dir):
        """Create an OutputFormatter instance for testing."""
        config = sample_config.copy()
        config['output_dir'] = temp_dir
        return OutputFormatter(config)

    def test_init_with_config(self, sample_config):
        """Test OutputFormatter initialization with configuration."""
        formatter = OutputFormatter(sample_config)
        
        assert formatter.config == sample_config
        assert formatter.output_dir == sample_config['output_dir']

    def test_init_with_defaults(self):
        """Test OutputFormatter initialization with minimal config."""
        minimal_config = {}
        formatter = OutputFormatter(minimal_config)
        
        assert formatter.output_dir == './output'

    def test_create_markdown_file(self, formatter, temp_dir):
        """Test Markdown file creation."""
        content = "# Test Title\n\nThis is test content."
        title = "Test Document"
        
        output_file = formatter.create_markdown(content, title, temp_dir)
        
        assert Path(output_file).exists()
        assert Path(output_file).suffix == '.md'
        
        with open(output_file, 'r', encoding='utf-8') as f:
            file_content = f.read()
        
        assert "# Test Title" in file_content
        assert "This is test content." in file_content

    def test_create_text_file(self, formatter, temp_dir):
        """Test plain text file creation."""
        content = "This is plain text content for testing."
        title = "Test Text Document"
        
        output_file = formatter.create_text(content, title, temp_dir)
        
        assert Path(output_file).exists()
        assert Path(output_file).suffix == '.txt'
        
        with open(output_file, 'r', encoding='utf-8') as f:
            file_content = f.read()
        
        assert content in file_content

    @patch('video_draft_creator.output_formatter.Document')
    def test_create_docx_file(self, mock_document_class, formatter, temp_dir):
        """Test DOCX file creation."""
        # Mock the Document class and its methods
        mock_doc = Mock()
        mock_document_class.return_value = mock_doc
        mock_doc.add_heading = Mock()
        mock_doc.add_paragraph = Mock()
        mock_doc.save = Mock()
        
        content = "This is test content for DOCX creation."
        title = "Test DOCX Document"
        
        output_file = formatter.create_docx(content, title, temp_dir)
        
        assert Path(output_file).suffix == '.docx'
        
        # Verify Document methods were called
        mock_doc.add_heading.assert_called()
        mock_doc.add_paragraph.assert_called()
        mock_doc.save.assert_called_with(output_file)

    def test_format_content_with_title(self, formatter):
        """Test content formatting with title."""
        content = "Sample content"
        title = "Sample Title"
        
        formatted = formatter._format_content_with_title(content, title)
        
        assert title in formatted
        assert content in formatted

    def test_format_content_with_metadata(self, formatter):
        """Test content formatting with metadata."""
        content = "Sample content"
        title = "Sample Title"
        metadata = {
            'duration': 120.5,
            'language': 'en',
            'created_at': '2024-01-01 12:00:00'
        }
        
        formatted = formatter._format_content_with_metadata(content, title, metadata)
        
        assert title in formatted
        assert content in formatted
        assert 'Duration: 2.0 minutes' in formatted
        assert 'Language: en' in formatted
        assert '2024-01-01 12:00:00' in formatted

    def test_generate_filename(self, formatter):
        """Test filename generation."""
        title = "Test Video Title"
        extension = "md"
        
        filename = formatter._generate_filename(title, extension)
        
        assert filename.endswith('.md')
        assert 'test' in filename.lower()
        assert 'video' in filename.lower()

    def test_generate_filename_with_special_chars(self, formatter):
        """Test filename generation with special characters."""
        title = "Test/Video\\Title:With*Special<Chars>?"
        extension = "txt"
        
        filename = formatter._generate_filename(title, extension)
        
        # Should remove or replace special characters
        assert '/' not in filename
        assert '\\' not in filename
        assert ':' not in filename
        assert '*' not in filename
        assert '<' not in filename
        assert '>' not in filename
        assert '?' not in filename

    def test_sanitize_filename(self, formatter):
        """Test filename sanitization."""
        unsafe_name = "unsafe/filename\\with:chars*"
        safe_name = formatter._sanitize_filename(unsafe_name)
        
        assert '/' not in safe_name
        assert '\\' not in safe_name
        assert ':' not in safe_name
        assert '*' not in safe_name

    def test_ensure_output_dir(self, formatter, temp_dir):
        """Test output directory creation."""
        new_dir = Path(temp_dir) / "new_output_dir"
        
        formatter._ensure_output_dir(str(new_dir))
        
        assert new_dir.exists()
        assert new_dir.is_dir()

    def test_create_all_formats(self, formatter, temp_dir):
        """Test creating all supported formats."""
        content = "# Test Content\n\nThis is test content for all formats."
        title = "Multi-Format Test"
        metadata = {'duration': 60.0, 'language': 'en'}
        
        with patch('video_draft_creator.output_formatter.Document') as mock_doc_class:
            mock_doc = Mock()
            mock_doc_class.return_value = mock_doc
            mock_doc.add_heading = Mock()
            mock_doc.add_paragraph = Mock()
            mock_doc.save = Mock()
            
            files = formatter.create_all_formats(content, title, temp_dir, metadata)
        
        assert len(files) == 3
        assert any(f.endswith('.md') for f in files)
        assert any(f.endswith('.txt') for f in files)
        assert any(f.endswith('.docx') for f in files)

    def test_format_duration(self, formatter):
        """Test duration formatting."""
        # Test seconds
        duration_str = formatter._format_duration(45.0)
        assert duration_str == "45.0 seconds"
        
        # Test minutes
        duration_str = formatter._format_duration(125.0)
        assert duration_str == "2.1 minutes"
        
        # Test hours
        duration_str = formatter._format_duration(3665.0)
        assert duration_str == "1.0 hours"

    def test_format_file_size(self, formatter):
        """Test file size formatting."""
        # Test bytes
        size_str = formatter._format_file_size(500)
        assert "500" in size_str and "bytes" in size_str
        
        # Test KB
        size_str = formatter._format_file_size(1500)
        assert "KB" in size_str
        
        # Test MB
        size_str = formatter._format_file_size(1500000)
        assert "MB" in size_str

    def test_create_summary_document(self, formatter, temp_dir):
        """Test summary document creation."""
        summary_data = {
            'title': 'Test Video',
            'original_text': 'Original transcript text',
            'corrected_text': 'Corrected transcript text',
            'summary': 'Brief summary',
            'keywords': ['test', 'video', 'summary'],
            'metadata': {
                'duration': 120.0,
                'language': 'en',
                'file_size': 1024000
            }
        }
        
        output_file = formatter.create_summary_document(summary_data, temp_dir)
        
        assert Path(output_file).exists()
        
        with open(output_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        assert 'Test Video' in content
        assert 'Brief summary' in content
        assert 'test, video, summary' in content or 'test\nvideo\nsummary' in content

    def test_create_transcript_with_timestamps(self, formatter, sample_transcript, temp_dir):
        """Test transcript creation with timestamps."""
        title = "Timestamped Transcript"
        
        output_file = formatter.create_transcript_with_timestamps(sample_transcript, title, temp_dir)
        
        assert Path(output_file).exists()
        
        with open(output_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        # Should contain timestamps and text
        assert '[0.0s]' in content or '00:00:00' in content
        assert 'Hello world' in content

    def test_create_structured_document(self, formatter, temp_dir):
        """Test structured document creation."""
        data = {
            'title': 'Structured Document Test',
            'sections': [
                {'heading': 'Introduction', 'content': 'This is the introduction.'},
                {'heading': 'Main Content', 'content': 'This is the main content.'},
                {'heading': 'Conclusion', 'content': 'This is the conclusion.'}
            ],
            'metadata': {
                'created_at': '2024-01-01',
                'language': 'en'
            }
        }
        
        output_file = formatter.create_structured_document(data, temp_dir)
        
        assert Path(output_file).exists()
        
        with open(output_file, 'r', encoding='utf-8') as f:
            content = f.read()
        
        assert 'Structured Document Test' in content
        assert 'Introduction' in content
        assert 'Main Content' in content
        assert 'Conclusion' in content

    def test_error_handling_invalid_output_dir(self, formatter):
        """Test error handling for invalid output directory."""
        content = "Test content"
        title = "Test"
        invalid_dir = "/invalid/nonexistent/path"
        
        # Should handle the error gracefully
        try:
            formatter.create_markdown(content, title, invalid_dir)
        except Exception as e:
            # Should be a meaningful error message
            assert 'directory' in str(e).lower() or 'path' in str(e).lower()

    def test_unicode_content_handling(self, formatter, temp_dir):
        """Test handling of Unicode content."""
        content = "测试内容 Test Content émojis 🎉 русский текст"
        title = "Unicode Test"
        
        # Test all formats with Unicode content
        md_file = formatter.create_markdown(content, title, temp_dir)
        txt_file = formatter.create_text(content, title, temp_dir)
        
        # Verify content is preserved
        with open(md_file, 'r', encoding='utf-8') as f:
            md_content = f.read()
        assert "测试内容" in md_content
        assert "émojis 🎉" in md_content
        assert "русский текст" in md_content
        
        with open(txt_file, 'r', encoding='utf-8') as f:
            txt_content = f.read()
        assert "测试内容" in txt_content

    def test_large_content_handling(self, formatter, temp_dir):
        """Test handling of large content."""
        # Create large content (simulate long transcript)
        large_content = "This is a test sentence. " * 10000  # ~250KB of text
        title = "Large Content Test"
        
        output_file = formatter.create_text(large_content, title, temp_dir)
        
        assert Path(output_file).exists()
        file_size = Path(output_file).stat().st_size
        assert file_size > 100000  # Should be large file


class TestCreateFormatterFromConfig:
    """Test cases for formatter factory function."""

    def test_create_from_dict_config(self, sample_config):
        """Test creating formatter from dictionary config."""
        formatter = create_formatter_from_config(sample_config)
        
        assert isinstance(formatter, OutputFormatter)
        assert formatter.config == sample_config

    def test_create_from_object_config(self, sample_config):
        """Test creating formatter from config object."""
        class MockConfig:
            def __init__(self, config_dict):
                for key, value in config_dict.items():
                    setattr(self, key, value)
        
        config_obj = MockConfig(sample_config)
        formatter = create_formatter_from_config(config_obj)
        
        assert isinstance(formatter, OutputFormatter)
        assert isinstance(formatter.config, dict)

    def test_create_with_none_config(self):
        """Test creating formatter with None config."""
        formatter = create_formatter_from_config(None)
        
        assert isinstance(formatter, OutputFormatter)
        assert formatter.config == {}


if __name__ == '__main__':
    pytest.main([__file__]) 